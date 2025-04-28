import os
import markdown
from docx import Document
from docx.shared import Pt, Inches
from docx2pdf import convert

def convert_markdown_to_pdf(input_file, output_file=None):
    """
    Markdownファイルを読み込み、PDFファイルに変換する
    
    Parameters:
        input_file (str): 入力となるMarkdownファイルのパス
        output_file (str, optional): 出力PDFファイルのパス。指定がない場合は入力ファイルの拡張子をpdfに変更したものを使用
    
    Returns:
        str: 出力されたPDFファイルのパス
    """
    # 出力ファイル名が指定されていない場合は、入力ファイル名の拡張子をpdfに変更
    if output_file is None:
        base_name = os.path.splitext(input_file)[0]
        output_file = f"{base_name}.pdf"
    
    # 中間DOCXファイルのパス
    docx_file = f"{os.path.splitext(input_file)[0]}_temp.docx"
    
    # Markdownファイルを読み込む
    try:
        with open(input_file, 'r', encoding='utf-8') as f:
            markdown_text = f.read()
    except FileNotFoundError:
        print(f"エラー: ファイル '{input_file}' が見つかりません。")
        return None
    except Exception as e:
        print(f"ファイル読み込み中にエラーが発生しました: {e}")
        return None
    
    # MarkdownをHTMLに変換
    html_text = markdown.markdown(
        markdown_text, 
        extensions=['tables', 'fenced_code', 'codehilite']
    )
    
    # DOCXドキュメントを作成
    doc = Document()
    
    # 基本スタイル設定
    style = doc.styles['Normal']
    style.font.name = 'Arial'
    style.font.size = Pt(11)
    
    # 余白設定
    sections = doc.sections
    for section in sections:
        section.top_margin = Inches(1)
        section.bottom_margin = Inches(1)
        section.left_margin = Inches(1)
        section.right_margin = Inches(1)
    
    # HTMLからテキストを抽出して追加（簡易的な実装）
    # 注: 実際のMarkdownの構造は保持されません。完全な変換にはより高度な処理が必要です。
    from bs4 import BeautifulSoup
    soup = BeautifulSoup(html_text, 'html.parser')
    
    # 見出しとテキストを追加
    for tag in soup.find_all(['h1', 'h2', 'h3', 'h4', 'h5', 'h6', 'p', 'pre', 'code', 'ul', 'ol', 'li']):
        if tag.name.startswith('h'):
            level = int(tag.name[1])
            heading = doc.add_heading(tag.get_text(), level=level)
        elif tag.name == 'p':
            doc.add_paragraph(tag.get_text())
        elif tag.name == 'pre' or tag.name == 'code':
            p = doc.add_paragraph(tag.get_text())
            p.style = 'No Spacing'
            for run in p.runs:
                run.font.name = 'Courier New'
        elif tag.name == 'ul' or tag.name == 'ol':
            # リストの処理は複雑なため、単純にテキストとして処理
            doc.add_paragraph(tag.get_text())
    
    # DOCXファイルを保存
    doc.save(docx_file)
    
    try:
        # DOCXをPDFに変換
        convert(docx_file, output_file)
        print(f"PDFが正常に生成されました: {output_file}")
        
        # 一時DOCXファイルを削除（必要に応じてコメントアウト）
        os.remove(docx_file)
        
        return output_file
    except Exception as e:
        print(f"PDF生成中にエラーが発生しました: {e}")
        return None

if __name__ == '__main__':
    # ここで直接入力ファイルと出力ファイルを指定
    input_file_path = "sample.md"  # ここに変換したいMDファイルのパスを指定
    output_file_path = None  # Noneにすると自動的に同名のPDFファイルを作成
    
    convert_markdown_to_pdf(input_file_path, output_file_path)